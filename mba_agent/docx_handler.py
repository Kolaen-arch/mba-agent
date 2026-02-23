"""
DOCX handler: read, write, and modify Word documents.
Uses python-docx for structured read/write operations.
"""

import os
import re
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def read_docx(filepath: str) -> dict:
    """
    Read a .docx file and extract structured content.
    Returns dict with full_text, sections (by heading), and metadata.
    """
    doc = Document(filepath)

    full_text_parts = []
    sections = []
    current_section = {"heading": "Preamble", "level": 0, "paragraphs": []}

    for para in doc.paragraphs:
        text = para.text.strip()
        style_name = para.style.name if para.style else ""

        # Detect headings
        if style_name.startswith("Heading"):
            # Save previous section
            if current_section["paragraphs"]:
                sections.append(current_section)

            try:
                level = int(style_name.replace("Heading ", "").replace("Heading", "0"))
            except ValueError:
                level = 1

            current_section = {"heading": text, "level": level, "paragraphs": []}
            full_text_parts.append(f"\n{'#' * level} {text}\n")
        elif text:
            current_section["paragraphs"].append(text)
            full_text_parts.append(text)

    # Don't forget last section
    if current_section["paragraphs"]:
        sections.append(current_section)

    # Extract tables
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        if table_data:
            tables.append(table_data)

    # Core properties
    props = doc.core_properties
    metadata = {
        "title": props.title or "",
        "author": props.author or "",
        "subject": props.subject or "",
        "created": str(props.created) if props.created else "",
        "modified": str(props.modified) if props.modified else "",
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "word_count": sum(len(p.text.split()) for p in doc.paragraphs),
        "page_estimate": sum(len(p.text.split()) for p in doc.paragraphs) // 250,
    }

    return {
        "full_text": "\n".join(full_text_parts),
        "sections": sections,
        "tables": tables,
        "metadata": metadata,
    }


def read_docx_section(filepath: str, heading: str) -> Optional[str]:
    """Read a specific section from a docx by heading name."""
    data = read_docx(filepath)
    for section in data["sections"]:
        if section["heading"].lower() == heading.lower():
            return "\n\n".join(section["paragraphs"])
    return None


def write_docx_from_markdown(
    markdown_text: str,
    output_path: str,
    title: str = "",
    author: str = "",
) -> str:
    """
    Convert markdown-formatted text to a .docx file.
    Handles headings (#, ##, ###), bold (**), italic (*), and paragraphs.
    """
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # Set properties
    if title:
        doc.core_properties.title = title
    if author:
        doc.core_properties.author = author

    lines = markdown_text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # Headings
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        else:
            # Regular paragraph — handle inline formatting
            para = doc.add_paragraph()
            _add_formatted_runs(para, line)

        i += 1

    doc.save(output_path)
    return output_path


def _add_formatted_runs(paragraph, text: str) -> None:
    """Parse inline markdown (bold, italic) into docx runs."""
    # Pattern: **bold**, *italic*, ***bold italic***
    pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*)'
    parts = re.split(pattern, text)

    i = 0
    while i < len(parts):
        part = parts[i]
        if part is None:
            i += 1
            continue
        if part.startswith("***") and part.endswith("***"):
            run = paragraph.add_run(parts[i + 1])
            run.bold = True
            run.italic = True
            i += 2
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(parts[i + 2] if i + 2 < len(parts) else part[2:-2])
            run.bold = True
            i += 3
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = paragraph.add_run(parts[i + 3] if i + 3 < len(parts) else part[1:-1])
            run.italic = True
            i += 4
        else:
            if part:
                paragraph.add_run(part)
            i += 1


def apply_changes_to_docx(
    source_path: str,
    changes: list[dict],
    output_path: str,
) -> str:
    """
    Apply a list of changes to a docx file.

    Each change is a dict:
    {
        "type": "replace" | "insert_after" | "delete" | "append",
        "section_heading": "...",      # for section-level ops
        "find_text": "...",            # for paragraph-level find
        "new_text": "...",             # replacement or new text
        "heading_level": 2,            # for insert_after with new heading
    }
    """
    doc = Document(source_path)
    changes_made = []

    for change in changes:
        ctype = change.get("type", "")

        if ctype == "replace" and change.get("find_text"):
            find = change["find_text"]
            replace = change.get("new_text", "")
            for para in doc.paragraphs:
                if find in para.text:
                    # Preserve formatting of first run, replace text
                    for run in para.runs:
                        if find in run.text:
                            run.text = run.text.replace(find, replace)
                            changes_made.append(f"Replaced: '{find[:50]}...' → '{replace[:50]}...'")
                            break

        elif ctype == "append" and change.get("new_text"):
            text = change["new_text"]
            if change.get("heading_level"):
                doc.add_heading(text, level=change["heading_level"])
                changes_made.append(f"Appended heading: {text[:50]}")
            else:
                para = doc.add_paragraph()
                _add_formatted_runs(para, text)
                changes_made.append(f"Appended paragraph: {text[:50]}...")

        elif ctype == "insert_after" and change.get("section_heading"):
            target_heading = change["section_heading"].lower()
            new_text = change.get("new_text", "")

            for idx, para in enumerate(doc.paragraphs):
                style_name = para.style.name if para.style else ""
                if style_name.startswith("Heading") and para.text.strip().lower() == target_heading:
                    # Find the end of this section (next heading of same or higher level)
                    insert_idx = idx + 1
                    try:
                        current_level = int(style_name.replace("Heading ", "").replace("Heading", "1"))
                    except ValueError:
                        current_level = 1

                    for j in range(idx + 1, len(doc.paragraphs)):
                        sn = doc.paragraphs[j].style.name if doc.paragraphs[j].style else ""
                        if sn.startswith("Heading"):
                            try:
                                jlevel = int(sn.replace("Heading ", "").replace("Heading", "1"))
                            except ValueError:
                                jlevel = 1
                            if jlevel <= current_level:
                                insert_idx = j
                                break
                    else:
                        insert_idx = len(doc.paragraphs)

                    # We can't easily insert at arbitrary positions with python-docx,
                    # so we append and note the location
                    para_new = doc.add_paragraph()
                    _add_formatted_runs(para_new, new_text)
                    changes_made.append(
                        f"Added after '{target_heading}': {new_text[:50]}... "
                        f"(note: appended at end — manual reorder may be needed)"
                    )
                    break

        elif ctype == "delete" and change.get("find_text"):
            find = change["find_text"]
            for para in doc.paragraphs:
                if find in para.text:
                    # Clear the paragraph (python-docx can't delete paragraphs easily)
                    for run in para.runs:
                        run.text = ""
                    para.text = ""
                    changes_made.append(f"Deleted paragraph containing: '{find[:50]}...'")
                    break

    doc.save(output_path)
    return "\n".join(changes_made) if changes_made else "No changes applied."


def export_full_paper(
    sections: list[dict],
    output_path: str,
    title: str = "",
    author: str = "",
    include_toc: bool = True,
) -> str:
    """
    Compile all paper sections into a single DOCX with APA 7th formatting.
    sections: list of dicts with {id, title, parent_id, content, order}.
    """
    doc = Document()

    # Set default font to Times New Roman 12pt
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # Set paragraph spacing (double-spaced for APA 7th)
    para_format = style.paragraph_format
    para_format.line_spacing = 2.0
    para_format.space_after = Pt(0)

    # Set margins to 1 inch
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title page
    if title:
        doc.core_properties.title = title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.space_before = Pt(120)
        run = title_para.add_run(title)
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "Times New Roman"

        if author:
            doc.core_properties.author = author
            author_para = doc.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_run = author_para.add_run(author)
            author_run.font.size = Pt(12)
            author_run.font.name = "Times New Roman"

        doc.add_page_break()

    # TOC placeholder
    if include_toc:
        toc_heading = doc.add_heading("Table of Contents", level=1)
        for sec in sorted(sections, key=lambda s: s.get("order", 0)):
            indent = "  " * (1 if sec.get("parent_id") else 0)
            doc.add_paragraph(f"{indent}{sec['id']} {sec['title']}")
        doc.add_page_break()

    # Sections
    for sec in sorted(sections, key=lambda s: s.get("order", 0)):
        # Determine heading level from parent_id
        level = 2 if sec.get("parent_id") else 1
        doc.add_heading(f"{sec['id']} {sec['title']}", level=level)

        content = sec.get("content", "")
        if content:
            for line in content.split("\n"):
                line = line.strip()
                if not line:
                    continue
                para = doc.add_paragraph()
                _add_formatted_runs(para, line)
        else:
            para = doc.add_paragraph()
            run = para.add_run("[Section content not yet written]")
            run.italic = True
            run.font.color.rgb = RGBColor(150, 150, 150)

    doc.save(output_path)
    return output_path


def compute_inline_diff(old_text: str, new_text: str) -> str:
    """Produce HTML with <ins> and <del> tags for word-level inline diff."""
    import difflib
    old_words = old_text.split()
    new_words = new_text.split()
    matcher = difflib.SequenceMatcher(None, old_words, new_words)
    parts = []
    for op, i1, i2, j1, j2 in matcher.get_opcodes():
        if op == 'equal':
            parts.append(' '.join(old_words[i1:i2]))
        elif op == 'replace':
            parts.append(f'<del>{" ".join(old_words[i1:i2])}</del>')
            parts.append(f'<ins>{" ".join(new_words[j1:j2])}</ins>')
        elif op == 'insert':
            parts.append(f'<ins>{" ".join(new_words[j1:j2])}</ins>')
        elif op == 'delete':
            parts.append(f'<del>{" ".join(old_words[i1:i2])}</del>')
    return ' '.join(parts)


def replace_section(source_path: str, heading: str, new_text: str, output_path: str) -> str:
    """Replace the content of a section (identified by heading) in a .docx file."""
    doc = Document(source_path)
    in_section = False
    section_level = 0
    paragraphs_to_clear = []

    for idx, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else ""

        if style_name.startswith("Heading"):
            try:
                level = int(style_name.replace("Heading ", "").replace("Heading", "1"))
            except ValueError:
                level = 1

            if in_section:
                # Hit next heading at same or higher level — stop
                if level <= section_level:
                    break
            elif para.text.strip().lower() == heading.strip().lower():
                in_section = True
                section_level = level
                continue
        elif in_section:
            paragraphs_to_clear.append(idx)

    if not paragraphs_to_clear:
        # Section not found or empty — append after heading
        doc.save(output_path)
        return "Section not found or empty — no changes made."

    # Clear old content
    for idx in paragraphs_to_clear:
        para = doc.paragraphs[idx]
        for run in para.runs:
            run.text = ""
        para.text = ""

    # Write new content into first cleared paragraph, rest as new paragraphs
    new_lines = [l for l in new_text.split("\n") if l.strip()]
    if new_lines and paragraphs_to_clear:
        first_para = doc.paragraphs[paragraphs_to_clear[0]]
        _add_formatted_runs(first_para, new_lines[0])
        for line in new_lines[1:]:
            new_para = doc.add_paragraph()
            _add_formatted_runs(new_para, line)

    doc.save(output_path)
    return f"Replaced section '{heading}' with {len(new_lines)} paragraphs."


def list_docx_files(directory: str) -> list[dict]:
    """List all .docx files in a directory with basic info."""
    path = Path(directory)
    files = []
    for f in sorted(path.glob("**/*.docx")):
        try:
            doc = Document(str(f))
            word_count = sum(len(p.text.split()) for p in doc.paragraphs)
            files.append({
                "path": str(f),
                "filename": f.name,
                "relative_path": str(f.relative_to(path)),
                "word_count": word_count,
                "page_estimate": word_count // 250,
            })
        except Exception:
            files.append({
                "path": str(f),
                "filename": f.name,
                "relative_path": str(f.relative_to(path)),
                "word_count": 0,
                "page_estimate": 0,
            })
    return files
